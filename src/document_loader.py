"""
文档加载器模块 — 支持多种格式文档的解析与加载

技术栈:
  - PyMuPDF (fitz):      PDF 文本、图片、元数据提取（快速，适合文本型 PDF）
  - pdfplumber:           PDF 表格精确提取（适合报表、数据型 PDF）
  - BeautifulSoup (bs4):  HTML/XML 解析与网页内容提取
  - python-docx:          Word 文档（.docx）解析

特性:
  - 统一入口: load(source) 根据扩展名自动路由
  - PDF 混合模式: PyMuPDF 提取文本 + pdfplumber 提取表格
  - 编码自动检测: 文本文件尝试 utf-8 → gbk → latin-1
  - 结构化元数据: 每个文档携带完整元数据信息
  - 密码保护 PDF 支持
  - 多格式支持: PDF / DOCX / TXT / MD / HTML / CSV / JSON / XML
  - 上下文管理器支持 (with 语句)
  - 从文件、字符串缓冲、URL 三种来源加载
"""

from __future__ import annotations

import csv
import io
import json
import os
import re
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Union

import fitz  # PyMuPDF
import pdfplumber
import requests
from bs4 import BeautifulSoup

try:
    from docx import Document as DocxDocument

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


# ---------------------------------------------------------------------------
# 自定义异常
# ---------------------------------------------------------------------------

class UnsupportedFormatError(ValueError):
    """不支持的文件格式"""

    def __init__(self, ext: str):
        super().__init__(f"不支持的文件格式: {ext}")
        self.ext = ext


class LoadError(RuntimeError):
    """文档加载失败"""


class FileSizeExceededError(LoadError):
    """文件大小超出限制"""


# ---------------------------------------------------------------------------
# DocumentLoader
# ---------------------------------------------------------------------------

class DocumentLoader:
    """多格式文档加载器

    用法:
        loader = DocumentLoader()
        result = loader.load("path/to/doc.pdf")
        result = loader.load("https://example.com")
        result = loader.load_string("<html>...</html>", ".html")

        # 上下文管理器
        with DocumentLoader() as loader:
            result = loader.load("path/to/doc.pdf")
    """

    # 扩展名 → 加载方法映射
    _EXT_LOADERS: Dict[str, str] = {
        ".pdf": "load_pdf",
        ".txt": "load_txt",
        ".md": "load_markdown",
        ".markdown": "load_markdown",
        ".doc": "load_doc",
        ".docx": "load_doc",
        ".html": "load_html",
        ".htm": "load_html",
        ".csv": "load_csv",
        ".json": "load_json",
        ".xml": "load_xml",
    }

    # 默认最大文件大小（字节）
    DEFAULT_MAX_FILE_SIZE: int = 500 * 1024 * 1024  # 500 MB

    def __init__(
        self,
        verbose: bool = False,
        max_file_size: Optional[int] = None,
    ):
        """
        Args:
            verbose:       是否打印加载日志
            max_file_size: 最大可加载文件大小（字节），None 使用默认值
        """
        self.verbose = verbose
        self.max_file_size = max_file_size or self.DEFAULT_MAX_FILE_SIZE
        self._loaded_count = 0

    def __repr__(self) -> str:
        return f"<DocumentLoader verbose={self.verbose} loaded={self._loaded_count}>"

    def __enter__(self) -> "DocumentLoader":
        return self

    def __exit__(self, *args) -> None:
        self._loaded_count = 0

    # ------------------------------------------------------------------
    # 文件校验与格式检测
    # ------------------------------------------------------------------

    def _check_file_size(self, file_path: str) -> None:
        """检查文件大小是否超出限制"""
        try:
            size = os.path.getsize(file_path)
            if size > self.max_file_size:
                raise FileSizeExceededError(
                    f"文件过大: {file_path} ({size / 1024 / 1024:.1f} MB) "
                    f"超出限制 {self.max_file_size / 1024 / 1024:.1f} MB"
                )
        except OSError as e:
            raise LoadError(f"无法访问文件: {file_path}\n  {e}") from e

    @staticmethod
    def _detect_format(file_path: str) -> Optional[str]:
        """通过文件内容魔数检测格式（无扩展名时使用）"""
        try:
            with open(file_path, "rb") as f:
                header = f.read(16)
            # PDF
            if header.startswith(b"%PDF"):
                return ".pdf"
            # XML / HTML
            if header.lstrip().startswith((b"<html", b"<!DOCTYPE", b"<?xml")):
                return ".xml"
            # ZIP-based formats (docx, xlsx)
            if header.startswith(b"PK"):
                return ".docx"
            # JSON (starts with { or [)
            if header.lstrip().startswith((b"{", b"[")):
                return ".json"
        except OSError:
            pass
        return None

    # ------------------------------------------------------------------
    # 从字符串/字节缓冲加载
    # ------------------------------------------------------------------

    def load_string(
        self,
        content: str,
        extension: str = ".txt",
        **kwargs,
    ) -> Dict[str, Any]:
        """从字符串内容加载文档

        Args:
            content:   文档内容字符串
            extension: 文件扩展名（决定解析方式），如 ``.html`` / ``.md``
            **kwargs:  传递给具体加载方法的参数

        Returns: 标准文档字典
        """
        ext = extension.lower() if extension.startswith(".") else f".{extension}"
        method_name = self._EXT_LOADERS.get(ext)
        if method_name is None:
            raise UnsupportedFormatError(ext)

        # 将字符串写入临时文件并加载
        import tempfile

        with tempfile.NamedTemporaryFile(
            mode="w", suffix=ext, delete=False, encoding="utf-8"
        ) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            loader = getattr(self, method_name)
            if self.verbose:
                print(f"[DocumentLoader] 从字符串加载 → {method_name}")
            return loader(tmp_path, **kwargs)
        finally:
            os.unlink(tmp_path)

    # ------------------------------------------------------------------
    # 统一入口
    # ------------------------------------------------------------------

    def load(self, source: Union[str, Path], **kwargs) -> Dict[str, Any]:
        """统一入口 — 根据输入自动选择加载方式

        Args:
            source: 文件路径 / URL / 扩展名（配合 content 参数使用）
            **kwargs: 传递给具体加载方法的参数
                - content: 从字符串内容加载（需同时指定 source 为扩展名）
                - password: PDF 解密密码

        Returns:
            {
                "content":   str,              # 完整文本
                "pages":     List[str],         # 按页/段分组
                "metadata":  Dict,              # 文档元数据
                "tables":    List[List],        # 表格（仅 PDF/DOCX）
                "method":    str,               # 使用的提取方法
            }
        """
        source = str(source)
        content = kwargs.pop("content", None)

        # ---- 从字符串内容加载 ----
        if content is not None:
            return self.load_string(content, source, **kwargs)

        # ---- URL ----
        if source.startswith(("http://", "https://")):
            return self.load_url(source, **kwargs)

        # ---- 文件 ----
        ext = Path(source).suffix.lower()

        # 无扩展名时尝试自动检测
        if not ext:
            detected = self._detect_format(source)
            if detected:
                ext = detected
                source = source + ext
            else:
                raise UnsupportedFormatError(source)

        # 文件大小校验
        self._check_file_size(source)

        method_name = self._EXT_LOADERS.get(ext)
        if method_name is None:
            raise UnsupportedFormatError(ext)

        loader = getattr(self, method_name)
        if self.verbose:
            print(f"[DocumentLoader] 加载 {source} → {method_name}")
        result = loader(source, **kwargs)
        self._loaded_count += 1
        return result

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def load_pdf(
        self,
        file_path: str,
        strategy: str = "hybrid",
        password: Optional[str] = None,
    ) -> Dict[str, Any]:
        """加载 PDF 文件

        Args:
            file_path:  PDF 路径
            strategy:   提取策略
                - ``"pymupdf"``    — 仅 PyMuPDF（快，适合纯文本 PDF）
                - ``"pdfplumber"`` — 仅 pdfplumber（精确，含表格）
                - ``"hybrid"``     — 混合（默认）：PyMuPDF 取文本 + pdfplumber 取表格
            password:   PDF 解密密码

        Returns: 标准文档字典
        """
        result: Dict[str, Any] = {
            "content": "",
            "pages": [],
            "metadata": {},
            "tables": [],
            "images": [],
            "method": strategy,
        }

        if strategy == "pymupdf":
            self._load_pdf_pymupdf(file_path, result, password=password)
        elif strategy == "pdfplumber":
            self._load_pdf_pdfplumber(file_path, result, password=password)
        else:  # hybrid
            self._load_pdf_pymupdf(file_path, result, password=password)
            if self._has_tables(file_path, password=password):
                tbl: Dict[str, Any] = {}
                self._load_pdf_pdfplumber(file_path, tbl, password=password)
                result["tables"] = tbl.get("tables", [])
                result["method"] = "hybrid"

        result["content"] = "\n\n".join(result["pages"])
        return result

    # -- PyMuPDF -------------------------------------------------------

    def _load_pdf_pymupdf(
        self,
        file_path: str,
        result: Dict[str, Any],
        password: Optional[str] = None,
    ) -> None:
        """使用 PyMuPDF 提取 PDF 全文、元数据与图片"""
        doc = fitz.open(file_path)
        # 尝试解密
        if password and doc.is_encrypted:
            if not doc.authenticate(password):
                raise LoadError("PDF 密码错误，无法解密文件")

        result["metadata"] = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "producer": doc.metadata.get("producer", ""),
            "page_count": len(doc),
            "file_size": os.path.getsize(file_path),
            "format": "pdf",
            "is_encrypted": doc.is_encrypted,
        }

        images_info = []
        for page_num, page in enumerate(doc):
            # 提取文本
            result["pages"].append(page.get_text())
            # 提取图片信息
            image_list = page.get_images(full=True)
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                images_info.append(
                    {
                        "page": page_num + 1,
                        "index": img_idx + 1,
                        "width": base_image.get("width", 0),
                        "height": base_image.get("height", 0),
                        "extension": base_image.get("ext", ""),
                        "size_bytes": len(base_image.get("image", b"")),
                    }
                )

        doc.close()
        result["images"] = images_info

    # -- pdfplumber ----------------------------------------------------

    def _load_pdf_pdfplumber(
        self,
        file_path: str,
        result: Dict[str, Any],
        password: Optional[str] = None,
    ) -> None:
        """使用 pdfplumber 提取 PDF（含表格检测）"""
        kwargs = {}
        if password:
            kwargs["password"] = password
        with pdfplumber.open(file_path, **kwargs) as pdf:
            result["metadata"]["page_count"] = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                result["pages"].append(text)
                tables = page.extract_tables()
                if tables:
                    result["tables"].extend(tables)

    # -- 表格检测 -------------------------------------------------------

    @staticmethod
    def _has_tables(file_path: str, password: Optional[str] = None) -> bool:
        """快速检测 PDF 前几页是否包含表格"""
        try:
            kwargs = {}
            if password:
                kwargs["password"] = password
            with pdfplumber.open(file_path, **kwargs) as pdf:
                return any(page.find_tables() for page in pdf.pages[:3])
        except Exception:
            return False

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def load_txt(
        self,
        file_path: str,
        encoding: Optional[str] = None,
    ) -> Dict[str, Any]:
        """加载纯文本文件（自动检测编码）

        Args:
            file_path: 文本文件路径
            encoding:  指定编码，为 ``None`` 时自动检测

        Returns: 标准文档字典
        """
        encodings: List[str]
        if encoding:
            encodings = [encoding]
        else:
            encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

        content = None
        used_encoding = None
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    content = f.read()
                used_encoding = enc
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if content is None:
            raise LoadError(
                f"无法解码文件（尝试了 {', '.join(encodings)}）: {file_path}"
            )

        lines = content.splitlines()
        return {
            "content": content,
            "pages": [content],
            "metadata": {
                "file_size": os.path.getsize(file_path),
                "encoding": used_encoding,
                "line_count": len(lines),
                "format": "txt",
            },
        }

    # ------------------------------------------------------------------
    # Markdown
    # ------------------------------------------------------------------

    def load_markdown(self, file_path: str) -> Dict[str, Any]:
        """加载 Markdown 文件，同时提取文档结构信息

        Returns: 标准文档字典 + metadata 中额外包含 headings / code_blocks
        """
        result = self.load_txt(file_path)
        content = result["content"]

        # 提取标题结构
        headings = re.findall(r"^(#{1,6})\s+(.+)$", content, re.MULTILINE)

        # 提取代码块
        code_blocks = re.findall(r"```(\w*)\n(.*?)```", content, re.DOTALL)

        result["metadata"].update(
            {
                "format": "markdown",
                "headings": [
                    {"level": len(m[0]), "text": m[1].strip()} for m in headings
                ],
                "code_blocks": len(code_blocks),
                "has_toc": bool(re.search(r"\[.*\]\(.*\)", content)),
            }
        )
        return result

    # ------------------------------------------------------------------
    # Word (.doc / .docx)
    # ------------------------------------------------------------------

    def load_doc(self, file_path: str) -> Dict[str, Any]:
        """加载 Word 文档（需要 python-docx）"""
        if not _HAS_DOCX:
            raise ImportError(
                "python-docx 未安装，请执行: pip install python-docx"
            )

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs]

        # 提取表格
        tables: List[List[List[str]]] = []
        for table in doc.tables:
            rows: List[List[str]] = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)

        content = "\n".join(paragraphs)
        return {
            "content": content,
            "pages": [content],
            "metadata": {
                "file_size": os.path.getsize(file_path),
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "format": "docx",
            },
            "tables": tables,
        }

    # ------------------------------------------------------------------
    # HTML
    # ------------------------------------------------------------------

    def load_html(self, file_path: str) -> Dict[str, Any]:
        """加载本地 HTML 文件"""
        result = self.load_txt(file_path)
        soup = BeautifulSoup(result["content"], "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        result["content"] = text
        result["pages"] = [text]
        result["metadata"].update(
            {
                "title": soup.title.string if soup.title else "",
                "format": "html",
            }
        )
        return result

    # ------------------------------------------------------------------
    # CSV
    # ------------------------------------------------------------------

    def load_csv(
        self,
        file_path: str,
        delimiter: str = ",",
        encoding: Optional[str] = None,
        max_rows: Optional[int] = None,
    ) -> Dict[str, Any]:
        """加载 CSV 文件

        Args:
            file_path:  CSV 文件路径
            delimiter:  分隔符（默认逗号，制表符用 ``\\t``）
            encoding:   文件编码
            max_rows:   最大读取行数（None 表示全部）

        Returns: 标准文档字典
        """
        txt_result = self.load_txt(file_path, encoding=encoding)
        used_encoding = txt_result["metadata"].get("encoding", "utf-8")

        rows: List[List[str]] = []
        with open(file_path, "r", encoding=used_encoding) as f:
            reader = csv.reader(f, delimiter=delimiter)
            for i, row in enumerate(reader):
                if max_rows and i > max_rows:
                    break
                rows.append(row)

        # 构建表格文本表示
        lines: List[str] = []
        for row in rows:
            lines.append(delimiter.join(row))
        content = "\n".join(lines)

        headers = rows[0] if rows else []
        return {
            "content": content,
            "pages": [content],
            "metadata": {
                "file_size": os.path.getsize(file_path),
                "format": "csv",
                "delimiter": delimiter,
                "encoding": used_encoding,
                "row_count": len(rows),
                "column_count": len(headers),
                "headers": headers,
            },
            "tables": rows,
        }

    # ------------------------------------------------------------------
    # JSON
    # ------------------------------------------------------------------

    def load_json(
        self,
        file_path: str,
        encoding: Optional[str] = None,
    ) -> Dict[str, Any]:
        """加载 JSON 文件，提取为可读文本

        Args:
            file_path: JSON 文件路径
            encoding:  文件编码

        Returns: 标准文档字典
        """
        txt_result = self.load_txt(file_path, encoding=encoding)
        used_encoding = txt_result["metadata"].get("encoding", "utf-8")

        with open(file_path, "r", encoding=used_encoding) as f:
            data = json.load(f)

        content = json.dumps(data, ensure_ascii=False, indent=2)

        return {
            "content": content,
            "pages": [content],
            "metadata": {
                "file_size": os.path.getsize(file_path),
                "format": "json",
                "encoding": used_encoding,
                "root_type": type(data).__name__,
                "key_count": len(data) if isinstance(data, dict) else None,
                "array_length": len(data) if isinstance(data, list) else None,
            },
        }

    # ------------------------------------------------------------------
    # XML
    # ------------------------------------------------------------------

    def load_xml(self, file_path: str, encoding: Optional[str] = None) -> Dict[str, Any]:
        """加载 XML 文件

        Args:
            file_path: XML 文件路径
            encoding:  文件编码

        Returns: 标准文档字典
        """
        txt_result = self.load_txt(file_path, encoding=encoding)
        used_encoding = txt_result["metadata"].get("encoding", "utf-8")

        tree = ET.parse(file_path)
        root = tree.getroot()

        # 提取文本内容
        text_parts: List[str] = []

        def _extract_text(element: ET.Element, level: int = 0) -> None:
            indent = "  " * level
            tag = element.tag.split("}")[-1]
            text = (element.text or "").strip()
            if text:
                text_parts.append(f"{indent}<{tag}>: {text}")
            elif list(element):
                text_parts.append(f"{indent}<{tag}>")
            else:
                text_parts.append(f"{indent}<{tag} />")
            for child in element:
                _extract_text(child, level + 1)

        _extract_text(root)

        content = "\n".join(text_parts)
        return {
            "content": content,
            "pages": [content],
            "metadata": {
                "file_size": os.path.getsize(file_path),
                "format": "xml",
                "encoding": used_encoding,
                "root_tag": root.tag.split("}")[-1],
                "num_elements": len(list(root.iter())),
                "attributes": list(root.attrib.keys()),
            },
        }

    # ------------------------------------------------------------------
    # URL
    # ------------------------------------------------------------------

    def load_url(
        self,
        url: str,
        timeout: int = 30,
        headers: Optional[Dict[str, str]] = None,
    ) -> Dict[str, Any]:
        """加载网页内容

        Args:
            url:     网页 URL
            timeout: 请求超时秒数
            headers: 自定义请求头

        Returns: 标准文档字典
        """
        default_headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            )
        }
        headers = {**default_headers, **(headers or {})}

        try:
            resp = requests.get(url, headers=headers, timeout=timeout)
            resp.raise_for_status()
            resp.encoding = resp.apparent_encoding
        except requests.RequestException as e:
            raise LoadError(f"无法访问 URL: {url}\n  {e}") from e

        soup = BeautifulSoup(resp.text, "html.parser")

        # 移除无用标签
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text)

        return {
            "content": text,
            "pages": [text],
            "metadata": {
                "url": url,
                "title": soup.title.string if soup.title else "",
                "status_code": resp.status_code,
                "content_length": len(resp.text),
                "format": "html",
            },
        }

    # ------------------------------------------------------------------
    # 辅助: 批量加载
    # ------------------------------------------------------------------

    def load_batch(
        self,
        sources: List[Union[str, Path]],
        **kwargs,
    ) -> List[Dict[str, Any]]:
        """批量加载多个文档"""
        return [self.load(s, **kwargs) for s in sources]


# ---------------------------------------------------------------------------
# 快捷入口
# ---------------------------------------------------------------------------

def load_document(source: Union[str, Path], **kwargs) -> Dict[str, Any]:
    """快捷函数 — 创建临时加载器并加载文档"""
    return DocumentLoader().load(source, **kwargs)


if __name__ == "__main__":
    loader = DocumentLoader()
    result = loader.load("测试文档.doc")
    print(result)